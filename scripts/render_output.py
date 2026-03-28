"""Utilities for rendering a markdown-first tailored resume to DOCX."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from utils import ensure_dir, sanitize_filename, timestamp_str


def _build_output_path(output_dir: Path, base_name: str, suffix: str) -> Path:
    """Construct a timestamped output path that does not overwrite the source file."""
    safe_name = sanitize_filename(base_name)
    return ensure_dir(output_dir) / f"{safe_name}_{timestamp_str()}{suffix}"


def _set_page_layout(document: Document) -> None:
    """Set compact resume-friendly page margins."""
    section = document.sections[0]
    section.orientation = WD_SECTION.CONTINUOUS
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def _set_default_style(document: Document) -> None:
    """Set a clean default style for all generated content."""
    style = document.styles["Normal"]
    style.font.name = "Garamond"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _remove_paragraph_spacing(paragraph) -> None:
    """Remove paragraph spacing for compact resume formatting."""
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0


def _set_tab_stop(paragraph, document: Document) -> None:
    """Add a right-aligned tab stop that matches the current page margins."""
    section = document.sections[0]
    usable_width_inches = (section.page_width - section.left_margin - section.right_margin) / 914400

    paragraph_properties = paragraph._p.get_or_add_pPr()
    tabs = paragraph_properties.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        paragraph_properties.append(tabs)

    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(usable_width_inches * 1440)))
    tabs.append(tab)


def _add_runs_with_bold_markers(paragraph, text: str) -> None:
    """Render inline markdown bold markers into DOCX runs."""
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        if part.startswith("**") and part.endswith("**"):
            run.bold = True


def _add_name(document: Document, text: str) -> None:
    """Render the name line."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _remove_paragraph_spacing(paragraph)
    run = paragraph.add_run(text.strip())
    run.bold = True
    run.font.name = "Garamond"
    run.font.size = Pt(25)


def _add_contact_line(document: Document, text: str) -> None:
    """Render the centered contact line."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _remove_paragraph_spacing(paragraph)
    run = paragraph.add_run(text.strip())
    run.font.name = "Garamond"
    run.font.size = Pt(11)


def _add_section_heading(document: Document, text: str) -> None:
    """Render a section heading with a bottom rule."""
    paragraph = document.add_paragraph()
    _remove_paragraph_spacing(paragraph)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(1)

    run = paragraph.add_run(text.strip())
    run.bold = True
    run.font.name = "Garamond"
    run.font.size = Pt(11)

    paragraph_properties = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    border.append(bottom)
    paragraph_properties.append(border)


def _add_two_column_bold_line(document: Document, text: str) -> None:
    """Render a bold left column and right-aligned second column."""
    paragraph = document.add_paragraph()
    _remove_paragraph_spacing(paragraph)
    paragraph.paragraph_format.space_before = Pt(3)
    _set_tab_stop(paragraph, document)

    left, _, right = text.partition("|")
    _add_runs_with_bold_markers(paragraph, left.strip())

    if right.strip():
        right_run = paragraph.add_run("\t" + right.strip())
        right_run.font.name = "Garamond"
        right_run.font.size = Pt(11)


def _add_project_title_line(document: Document, text: str) -> None:
    """Render a project line with inline tools instead of right-tab alignment."""
    paragraph = document.add_paragraph()
    _remove_paragraph_spacing(paragraph)
    paragraph.paragraph_format.space_before = Pt(3)

    left, _, right = text.partition("|")
    name = left.strip().strip("*")

    run = paragraph.add_run(name)
    run.bold = True
    run.font.name = "Garamond"
    run.font.size = Pt(11)

    if right.strip():
        tools_run = paragraph.add_run(" | " + right.strip())
        tools_run.font.name = "Garamond"
        tools_run.font.size = Pt(11)


def _add_subline_tab(document: Document, text: str) -> None:
    """Render an italic left column and right-aligned second column."""
    paragraph = document.add_paragraph()
    _remove_paragraph_spacing(paragraph)
    _set_tab_stop(paragraph, document)

    left, _, right = text.partition("|")
    left_clean = left.strip().strip("*")

    run = paragraph.add_run(left_clean)
    run.italic = True
    run.font.name = "Garamond"
    run.font.size = Pt(11)

    if right.strip():
        right_run = paragraph.add_run("\t" + right.strip())
        right_run.font.name = "Garamond"
        right_run.font.size = Pt(11)


def _add_subline(document: Document, text: str) -> None:
    """Render a standalone italic line."""
    paragraph = document.add_paragraph()
    _remove_paragraph_spacing(paragraph)
    run = paragraph.add_run(text.strip())
    run.italic = True
    run.font.name = "Garamond"
    run.font.size = Pt(11)


def _add_bullet(document: Document, text: str) -> None:
    """Render a compact bullet point with inline bold markers."""
    paragraph = document.add_paragraph(style="List Bullet")
    _remove_paragraph_spacing(paragraph)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    _add_runs_with_bold_markers(paragraph, text.strip())


def _add_body_line(document: Document, text: str) -> None:
    """Render a plain body line with inline bold markers."""
    paragraph = document.add_paragraph()
    _remove_paragraph_spacing(paragraph)
    _add_runs_with_bold_markers(paragraph, text.strip())


def _render_markdown_to_docx(markdown_content: str, output_path: Path) -> None:
    """Convert strict resume markdown into a styled DOCX file."""
    lines = markdown_content.splitlines()
    document = Document()
    _set_page_layout(document)
    _set_default_style(document)

    first_heading_done = False
    current_section = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# ") and not first_heading_done:
            _add_name(document, stripped[2:])
            first_heading_done = True
        elif first_heading_done and not stripped.startswith("## ") and "|" in stripped and "@" in stripped:
            _add_contact_line(document, stripped)
        elif stripped.startswith("## "):
            current_section = stripped[3:].upper()
            _add_section_heading(document, stripped[3:])
        elif stripped.startswith("- "):
            _add_bullet(document, stripped[2:])
        elif stripped.startswith("**") and "|" in stripped:
            if current_section == "PROJECTS":
                _add_project_title_line(document, stripped)
            else:
                _add_two_column_bold_line(document, stripped)
        elif stripped.startswith("*") and "|" in stripped:
            if current_section == "PROJECTS":
                continue
            _add_subline_tab(document, stripped)
        elif stripped.startswith("*") and stripped.endswith("*"):
            _add_subline(document, stripped.strip("*"))
        else:
            _add_body_line(document, stripped)

    document.save(output_path)


def save_docx_output(output_dir: Path, base_name: str, content: str) -> Path:
    """Render markdown resume content to a styled DOCX file."""
    output_path = _build_output_path(output_dir, base_name, ".docx")
    _render_markdown_to_docx(content.strip(), output_path)
    return output_path


def main() -> None:
    """Render markdown resume content to a DOCX output file."""
    parser = argparse.ArgumentParser(description="Render strict resume markdown to a DOCX file.")
    parser.add_argument("--content", help="Inline markdown resume content.")
    parser.add_argument("--content-file", help="Path to a file containing markdown resume content.")
    parser.add_argument(
        "--output-dir",
        required=True,
        help="Directory where the rendered output file should be saved.",
    )
    parser.add_argument(
        "--base-name",
        default="tailored_resume",
        help="Base file name for the rendered output.",
    )
    args = parser.parse_args()

    if not args.content and not args.content_file:
        raise ValueError("Provide either --content or --content-file.")
    if args.content and args.content_file:
        raise ValueError("Provide only one of --content or --content-file.")

    if args.content_file:
        content = Path(args.content_file).read_text(encoding="utf-8")
    else:
        content = args.content

    output_dir = Path(args.output_dir)
    output_path = save_docx_output(output_dir, args.base_name, content)

    print(json.dumps({"output_file": str(output_path)}, indent=2))


if __name__ == "__main__":
    main()
